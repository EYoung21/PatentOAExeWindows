# How to build exe if don't already have it
# 1. You must have pytesseract (https://github.com/UB-Mannheim/tesseract/wiki) downlaoded and located in the default folder: C:\Program Files\Tesseract-OCR\tesseract.exe
# 2. run: "pyinstaller oa_processor.spec" in the terminal
# 3. right click the "dist" folder
# 4. click "reveal in file explorer"
# 5. open the dist folder, and copy (ctrl c) the exe
# 6. paste the exe you want to run it in and double click it to run it

import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
import os
import re
import csv
from dotenv import load_dotenv
from openai import OpenAI
import sys
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from selenium.common.exceptions import TimeoutException
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#(os.path.join(tesseract_path, 'leptonica-1.82.0.dll'), '.'),
def get_tesseract_path():
    if getattr(sys, 'frozen', False):
        # If the application is run as a bundle, use the path relative to the executable
        return os.path.join(sys._MEIPASS, 'tesseract.exe')
    else:
        # If it's run as a normal Python script, use the default Tesseract path
        return r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set the Tesseract command
pytesseract.pytesseract.tesseract_cmd = get_tesseract_path()

def get_env_path():
    if getattr(sys, 'frozen', False):
        # If the application is run as a bundle, the .env file 
        # is located in the same directory as the executable
        return os.path.join(sys._MEIPASS, '.env')
    else:
        # If it's run as a normal Python script, use the current directory
        return '.env'
    
def create_docx_summary(subfolder_path, total_pdf_summary, obj):
    doc = Document()
    
    # Add a title
    title = doc.add_heading(f'Summary for Application {obj.applicationID}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add application details
    details = doc.add_paragraph()
    details.add_run('Application Details:').bold = True
    details.add_run(f'\nApplication ID: {obj.applicationID}')
    details.add_run(f'\nReference Number: {obj.refrenceNumber}')
    details.add_run(f'\nDue Date: {obj.dueDate}')
    details.add_run(f'\nExaminer: {obj.examinerName}')
    details.add_run(f'\nExaminer Phone: {str(obj.phone_numbers[0])[22:] if obj.phone_numbers else "N/A"}')

    # Add a line separator
    doc.add_paragraph('_' * 50)

    # Split the summary into sections
    sections = total_pdf_summary.split('\n\n')

    for section in sections:
        if section.strip():
            # Add section title
            if 'Summary of' in section:
                doc.add_heading(section.split('\n')[0], level=2)
                content = '\n'.join(section.split('\n')[1:])
            else:
                content = section

            # Add section content
            paragraph = doc.add_paragraph(content)

            # Add some space after each section
            paragraph.add_run().add_break()

    # Save the document
    doc_path = os.path.join(subfolder_path, f'{obj.applicationID.replace('/', '_')}_summary.docx')
    doc.save(doc_path)
    print(f"Summary saved as {doc_path}")

# Load the .env file
load_dotenv(dotenv_path=get_env_path())

class Solution():

    def __init__(self):
        pass

    # def defineUserInput(self):
    #     self.user_input = input("Which file would you like to examine? ")
    #     return self.user_input
    
    def inputRefReturnText(self, ref, subfolder_path):
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.common.exceptions import TimeoutException
        import requests
        import os

        driver = webdriver.Chrome()  # Use the appropriate driver for your browser
        driver.get('https://patents.google.com/')

        wait = WebDriverWait(driver, 30)

        try:
            search_bar = wait.until(EC.presence_of_element_located((By.NAME, "q")))
            initial_url = driver.current_url
            search_bar.send_keys(ref)
            search_bar.send_keys(Keys.RETURN)

            wait.until(lambda driver: driver.current_url != initial_url)
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '#wrapper > div:nth-child(3)')))

            # Look for the PDF download link
            pdf_link = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "a.style-scope.patent-result[target='_blank']")))
            pdf_url = pdf_link.get_attribute('href')

            # Download the PDF
            response = requests.get(pdf_url)
            if response.status_code == 200:
                pdf_filename = f"{ref.replace('/', '_')} PDF.pdf"  # Replace '/' with '_' in filename
                pdf_path = os.path.join(subfolder_path, pdf_filename)
                with open(pdf_path, 'wb') as f:
                    f.write(response.content)
                print(f"PDF downloaded and saved as {pdf_filename}")
            else:
                print(f"Failed to download PDF for {ref}")

            # Get the page text
            html = driver.page_source
            soup = BeautifulSoup(html, 'html.parser')
            page_text = soup.get_text(separator='\n', strip=True)

        except TimeoutException:
            print(f"Timed out waiting for search results or PDF link to load for {ref}.")
            driver.quit()
            return None
        except Exception as e:
            print(f"An error occurred while processing {ref}: {str(e)}")
            driver.quit()
            return None

        driver.quit()
        return page_text[:-len(page_text)//5]

    def extract_text_from_pdf(self, PDFPath):
        
        pdf_document = fitz.open(PDFPath)

        extracted_text = ""
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            # Extract text
            text = page.get_text()
            extracted_text += text
            
            # If no text found, look for images
            if not text.strip():  # No text found
                # print(f"No text found on page {page_num + 1}. Now searching for images.")
                image_list = page.get_images(full=True)

                if image_list:
                    for img_index, img in enumerate(image_list, start=1):
                        # Extract the image bytes
                        base_image = pdf_document.extract_image(img[0])
                        image_bytes = base_image["image"]

                        # Convert to a PIL Image
                        image = Image.open(io.BytesIO(image_bytes))

                        # Apply OCR to the image
                        ocr_text = pytesseract.image_to_string(image)
                        extracted_text += ocr_text
                        # print(f"Text extracted from image {img_index} on page {page_num + 1}.")

        return extracted_text
    
    def defineREGEX(self, stringInput):
        # Normalize the text, preserving newlines
        normalized_text = re.sub(r'[^\S\n]+', ' ', stringInput)

        # Define patterns
        application_id_pattern = r"\d{2}/\d{3},\d{3}"
        reference_number_pattern = r"\d{4}-[A-Za-z0-9]+"
        due_date_pattern = r"\d{2}/\d{2}/\d{4}"
        examiner_name_pattern = r"[A-Z]+,\s*[A-Z]+(?:\s[A-Z]\.?)?"

        # Updated telephone pattern to handle line breaks
        updated_telephone_pattern = r"whose\s+telephone\s+number\s+is\s*(?:\(?\d{3}\)?[-.\s]?\n?){2}\d{4}"
#11,995,475
        ref_pulled_pattern = r"\d{4}/\d{7}|\d{2},\d{3},\d{3}|\d{1},\d{3},\d{3}"

        typePattern = r"THIS ACTION IS MADE FINAL"
        # ref_pulled_pattern2 = 

        # Find patterns
        application_id = re.findall(application_id_pattern, normalized_text)
        phones = re.findall(updated_telephone_pattern, normalized_text, re.DOTALL)
        refMatches = re.findall(reference_number_pattern, normalized_text)
        dateMatches = re.findall(due_date_pattern, normalized_text)
        examiner_name_match = re.findall(examiner_name_pattern, stringInput)

        pulledRefMatches = re.findall(ref_pulled_pattern, normalized_text)

        typeMatches = re.findall(typePattern, normalized_text)

        # Assign values
        self.applicationID = application_id[0] if application_id else None
        self.refrenceNumber = refMatches[1] if len(refMatches) > 1 else None
        self.dueDate = dateMatches[1] if len(dateMatches) > 1 else None
        self.examinerName = examiner_name_match[0] if examiner_name_match else None
        self.phone_numbers = [re.sub(r'\s+', '', phone) for phone in phones]  # Clean up extracted phone numbers
        self.totalFinalTypes = typeMatches

        self.total_pulled_refs = pulledRefMatches if pulledRefMatches else None

#pyinstaller oa_processor.spec  


import traceback
import sys

def main():
    try:
        directory = os.getcwd()
        print(f"Processing files in: {directory}")

        # Check if the API key is set
        if not os.getenv("OPENAI_API_KEY"):
            print("Error: OPENAI_API_KEY is not set in the .env file.")
            return

        directory_name = os.path.join(directory, "DataCSV, Refrence Summaries & PDFs")

        # Create the directory
        try:
            os.makedirs(directory_name, exist_ok=True)
            print(f"Directory '{directory_name}' created successfully.")
        except PermissionError:
            print(f"Permission denied: Unable to create directory '{directory_name}'. Please check your permissions.")
            return
        except Exception as e:
            print(f"An error occurred while creating directory: {e}")
            return

        # Define the CSV file path
        csv_file_path = os.path.join(directory_name, "data.csv")
        header = ["Appl. No", "Ref. No.", "Due Date", "Due Date", "IHC", "G", "Examiner's Name", "Examiner's Phone Number"]
        rows = []

        for filename in os.listdir(directory):
            if filename.lower().endswith('.pdf'):
                f = os.path.join(directory, filename)
                print(f"Processing file: {filename}")
                if os.path.isfile(f):
                    subfolder_name = f"MATERIALS FOR {filename}"
                    subfolder_path = os.path.join(directory_name, subfolder_name)

                    try:
                        os.makedirs(subfolder_path, exist_ok=True)
                        print(f"Directory '{subfolder_path}' created successfully.")
                    except PermissionError:
                        print(f"Permission denied: Unable to create subfolder '{subfolder_path}'. Skipping this file.")
                        continue
                    except Exception as e:
                        print(f"An error occurred while creating subfolder: {e}")
                        continue

                    try:
                        obj = Solution()
                        text = obj.extract_text_from_pdf(f)
                        obj.defineREGEX(text)

                        # Initialize variables with default values
                        applicationID = obj.applicationID if obj.applicationID else "N/A"
                        refrenceNumber = obj.refrenceNumber if obj.refrenceNumber else "N/A"
                        dueDate = obj.dueDate if obj.dueDate else "N/A"
                        examinerRealName = "N/A"
                        examinerNumber = "N/A"

                        print(f"ID: {applicationID}")
                        print(f"Refrence #: {refrenceNumber}")
                        print(f"DueDate: {dueDate}")

                        if obj.examinerName:
                            arr = obj.examinerName.split(",")
                            examinerRealName = f"{arr[1].strip()} {arr[0].strip()}" if len(arr) > 1 else obj.examinerName
                        print(f"Examiner name: {examinerRealName}")

                        if obj.phone_numbers:
                            examinerNumber = str(obj.phone_numbers[0])[22:] if obj.phone_numbers else "N/A"
                        print(f"Examiner #: {examinerNumber}")

                        if obj.total_pulled_refs:
                            print(f"Pulled refs: {obj.total_pulled_refs}")
                        else:
                            obj.total_pulled_refs = []
                            print("Pulled refs not found")

                        print(f"OA Type? {obj.totalFinalTypes}")

                        oaType = "FOA Response 2mo" if obj.totalFinalTypes else "OA Response"

                        rows.append([applicationID, refrenceNumber, oaType, dueDate, "", "", examinerRealName, examinerNumber])

                        #here generate summary and put in subfolder_path

                        total_pdf_summary = ""

                        # prompt = f"Summarize this document:\n\n{text}"

                        prompt = f"Can you summarize this technology including telling me what is the problem it's trying to solve, a detailed summary of the technology and keywords and their definitions used?\n\n{text}"

                        # load_dotenv()

                        client = OpenAI(
                            # This is the default and can be omitted
                            # api_key=os.environ.get("OPENAI_API_KEY"),
                            api_key = os.getenv("OPENAI_API_KEY")
                        )
                        #stuff to summarize OAPDF
                        # # print(os.getenv("OPENAI_API_KEY"))
                        # response = client.chat.completions.create(
                        #     model="gpt-4o",
                        #     messages=[
                        #         {
                        #             "role": "user",
                        #             "content": [
                        #                 {"type": "text", "text": prompt}, #can individualize prompts later
                        #                 # {
                        #                 #     "type": "image_url",
                        #                 #     "image_url": {"url": f"{img_url}"},
                        #                 # },
                        #             ],
                        #         }
                        #     ],
                        # )

                        # content = response.choices[0].message.content

                        # total_pdf_summary += "Summary of original office action({OAName}):"
                        # total_pdf_summary += "\n"
                        # total_pdf_summary += content
                        # total_pdf_summary += "\n"
                        # total_pdf_summary += "\n"
                        
                        applicationIDPatenttext = obj.inputRefReturnText(obj.applicationID, subfolder_path)
                        if applicationIDPatenttext is not None:
                            prompt2 = f"Can you summarize this technology including telling me what is the problem it's trying to solve, a detailed summary of the technology and keywords and their definitions used?\n\n{applicationIDPatenttext}"
                            response2 = client.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[
                                    {
                                        "role": "user",
                                        "content": [
                                            {"type": "text", "text": prompt2}, #can individualize prompts later
                                            # {
                                            #     "type": "image_url",
                                            #     "image_url": {"url": f"{img_url}"},
                                            # },
                                        ],
                                    }
                                ],
                            )
                            content2 = response2.choices[0].message.content
                            total_pdf_summary += f"Summary of application ID({obj.applicationID})'s google patent text"
                            total_pdf_summary += "\n"
                            total_pdf_summary += content2
                            total_pdf_summary += "\n"
                            total_pdf_summary += "\n"
                        else:
                            print(f"Unable to fetch patent text for application ID: {obj.applicationID}")
                            # Handle this case appropriately, maybe skip this summary or use a placeholder



                        for i in range(len(obj.total_pulled_refs)):
                            refPatentText = obj.inputRefReturnText(obj.total_pulled_refs[i], subfolder_path)
                            if refPatentText is not None:
                                prompt = f"Can you summarize this technology including telling me what is the problem it's trying to solve, a detailed summary of the technology and keywords and their definitions used?:\n\n{refPatentText}"
                                response = client.chat.completions.create(
                                    model="gpt-4o-mini",
                                    messages=[
                                        {
                                            "role": "user",
                                            "content": [
                                                {"type": "text", "text": prompt}, #can individualize prompts later
                                            ],
                                        }
                                    ],
                                )
                                content = response.choices[0].message.content
                                total_pdf_summary += f"Summary of current refrence({obj.total_pulled_refs[i]})'s google patent text"
                                total_pdf_summary += "\n"
                                total_pdf_summary += content
                                total_pdf_summary += "\n"
                                total_pdf_summary += "\n"
                            else:
                                print(f"Unable to fetch patent text for reference: {obj.total_pulled_refs[i]}")


                        # with open(subfolder_path + '/summary.txt', 'w') as summary_file:
                        #     summary_file.write(str(total_pdf_summary))

                        create_docx_summary(subfolder_path, total_pdf_summary, obj)

                        # exit()
                    except Exception as e:
                        print(f"An error occurred while processing file {filename}: {str(e)}")
                        print(f"Skipping this file and continuing with the next...")
                        continue
            else:
                print(f"Skipping non-PDF file: {filename}")

        # Create the CSV file and write the header and rows
        try:
            with open(csv_file_path, mode='w', newline='') as file:
                writer = csv.writer(file)
                writer.writerow(header)
                writer.writerows(rows)
            print(f"CSV file '{csv_file_path}' created successfully with header and rows.")
        except PermissionError:
            print(f"Permission denied: Unable to create CSV file '{csv_file_path}'. Please check your permissions.")
        except Exception as e:
            print(f"An error occurred while creating the CSV file: {e}")

    except Exception as e:
        print(f"An error occurred: {str(e)}")
        print("Traceback:")
        traceback.print_exc()
    
    finally:
        # This will keep the console window open
        input("Press Enter to exit...")

if __name__ == "__main__":
    main()